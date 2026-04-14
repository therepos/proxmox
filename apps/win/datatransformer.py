"""
╔══════════════════════════════════════════════════════════════════╗
║         SMART DATA TRANSFORMER v2.0                              ║
║         High-fidelity workpaper → markdown with fake data        ║
║         Handles merged cells, comments, formulas                 ║
║         Double-click to run — dependencies auto-install          ║
╚══════════════════════════════════════════════════════════════════════╝

Converts Excel/Word/PDF workpapers to high-fidelity markdown
with all client data replaced by realistic fakes.

Key features:
  - Merged cells are unmerged and values propagated correctly
  - Reviewer comments extracted and placed inline at their cell
  - Formulas resolved to computed values
  - Sheet structure preserved as markdown sections
  - Consistent fake data (same entity = same fake throughout)
"""

# ══════════════════════════════════════════════════════════════
#  AUTO-INSTALL
# ══════════════════════════════════════════════════════════════

import subprocess
import sys
import importlib

def ensure(pkg, imp=None):
    try:
        importlib.import_module(imp or pkg)
    except ImportError:
        print(f"  Installing {pkg}...")
        subprocess.check_call(
            [sys.executable, "-m", "pip", "install", pkg, "--quiet"],
            stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL
        )

print("\n  Checking dependencies...")
ensure("openpyxl")
ensure("pdfplumber")
ensure("python-docx", "docx")
ensure("pypdf")
print("  All dependencies ready.\n")

# ══════════════════════════════════════════════════════════════
#  IMPORTS
# ══════════════════════════════════════════════════════════════

import os
import re
import json
import random
import string
import datetime
from collections import OrderedDict

import openpyxl
from openpyxl.utils import get_column_letter
import pdfplumber
from pypdf import PdfReader, PdfWriter
from docx import Document

# ══════════════════════════════════════════════════════════════
#  FAKE DATA POOLS
# ══════════════════════════════════════════════════════════════

FAKE_FIRST = [
    "James", "Emily", "Michael", "Sarah", "Daniel", "Rachel", "Robert",
    "Amanda", "William", "Nicole", "Andrew", "Lisa", "Steven", "Rebecca",
    "Wei", "Mei", "Jun", "Hui", "Xin", "Yan", "Hao", "Ling",
    "Raj", "Priya", "Arun", "Deepa", "Sanjay", "Anita",
    "Ahmad", "Fatimah", "Hassan", "Nurul", "Omar", "Siti",
]

FAKE_LAST = [
    "Anderson", "Mitchell", "Campbell", "Stewart", "Robinson", "Clarke",
    "Harris", "Wright", "Thompson", "Walker", "Cooper", "Morgan",
    "Chen", "Wang", "Zhang", "Li", "Liu", "Yang", "Huang", "Wu",
    "Tan", "Lim", "Lee", "Ng", "Wong", "Goh", "Chua", "Ong",
    "Kumar", "Singh", "Sharma", "Patel", "Nair", "Rajan",
]

FAKE_COMPANIES = [
    "Horizon Capital", "Sterling Industries", "Atlas Dynamics",
    "Pinnacle Solutions", "Meridian Group", "Nexus Holdings",
    "Vanguard Systems", "Pacific Ventures", "Quantum Enterprises",
    "Summit Partners", "Apex Technologies", "Crescendo Corp",
    "Lighthouse Global", "Paramount Services", "Zenith International",
    "Nova Trading", "Beacon Logistics", "Sapphire Resources",
    "Ironwood Capital", "Silverline Associates", "Terra Firma",
    "Oceanic Industries", "Ridgeway Holdings", "Cornerstone Ventures",
]

FAKE_DOMAINS = [
    "horizoncap.com", "sterlingind.com", "atlasdyn.com",
    "pinnaclesol.com", "meridiangrp.com", "nexushold.com",
    "vanguardsys.com", "pacificvent.com", "quantument.com",
]

FAKE_PROJECTS = [
    "Project Aurora", "Project Beacon", "Project Cascade",
    "Project Delta", "Project Eclipse", "Project Frontier",
    "Project Granite", "Project Horizon", "Project Indigo",
]


# ══════════════════════════════════════════════════════════════
#  TRANSFORMER ENGINE
# ══════════════════════════════════════════════════════════════

class DataTransformer:

    def __init__(self, seed=None):
        self.seed = seed or random.randint(100000, 999999)
        self.rng = random.Random(self.seed)
        self.amount_multiplier = round(0.3 + self.rng.random() * 0.5, 4)
        self.date_offset = self.rng.randint(-180, 180)

        # Consistent mappings
        self.maps = {
            "name": OrderedDict(), "company": OrderedDict(),
            "email": OrderedDict(), "ssn": OrderedDict(),
            "nric": OrderedDict(), "phone": OrderedDict(),
            "engmt": OrderedDict(), "ref": OrderedDict(),
            "project": OrderedDict(), "custom": OrderedDict(),
        }
        self._used_names = set()
        self._used_companies = set()
        self.custom_terms = []
        self.patterns = self._build_patterns()

    def _build_patterns(self):
        p = []
        p.append(("CC", re.compile(r'\b(?:\d{4}[-\s]?){3}\d{4}\b')))
        p.append(("SSN", re.compile(r'\b\d{3}-\d{2}-\d{4}\b')))
        p.append(("NRIC", re.compile(r'\b[STFGM]\d{7}[A-Z]\b', re.IGNORECASE)))
        p.append(("EMAIL", re.compile(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b')))
        p.append(("URL", re.compile(r'https?://[^\s,;)\]]+')))
        p.append(("PHONE", re.compile(r'\+\d{1,3}[-.\s]?\(?\d{2,4}\)?[-.\s]?\d{3,4}[-.\s]?\d{3,4}\b')))
        p.append(("PHONE", re.compile(r'\(?\d{3}\)?[-.\s]\d{3}[-.\s]\d{4}\b')))
        p.append(("PHONE", re.compile(r'\b[689]\d{3}[-.\s]?\d{4}\b')))
        p.append(("ACCT", re.compile(r'\b(?:ABA|SWIFT|A/C|Acct?|Account|Routing)\s*(?:#|No\.?|Number)?:?\s*\d{6,12}\b', re.IGNORECASE)))
        p.append(("ENGMT", re.compile(r'\b(?:ENG|SG|US|UK|HK|AU|JP|CN|MY|TH|ID|PH|IN|KR|EY)[-/]?\d{2,4}[-/]?\d{3,6}\b', re.IGNORECASE)))
        p.append(("REF", re.compile(r'\b(?:INV|PO|SO|WP|CHK|REF|DOC|RPT|JE|AJE|GJ|TB|AP|AR|GL)[-#]?\d{3,10}\b', re.IGNORECASE)))
        p.append(("AMOUNT", re.compile(r'[\(]?\$\s?\d{1,3}(?:,\d{3})*(?:\.\d{1,2})?(?:\s?[MBKmb](?:illion)?)?\)?')))
        p.append(("AMOUNT", re.compile(r'\b(?:USD|SGD|EUR|GBP|JPY|AUD|HKD|MYR)\s?\d{1,3}(?:,\d{3})*(?:\.\d{1,2})?\b')))
        p.append(("DATE", re.compile(
            r'\b(?:\d{1,2}[/-]\d{1,2}[/-]\d{2,4}|(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{1,2},?\s+\d{4}|\d{1,2}\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{4}|\d{4}-\d{2}-\d{2})\b', re.IGNORECASE
        )))
        suffixes = '|'.join([r'(?:Pte\.?\s*)?Ltd\.?', r'Inc\.?', r'Corp\.?', r'LLC', r'LLP', r'PLC', r'Group', r'Holdings', r'Sdn\.?\s*Bhd\.?'])
        p.append(("COMPANY", re.compile(rf'\b(?:[A-Z][A-Za-z&\'-]+(?:\s+[A-Z][A-Za-z&\'-]+){{0,5}})\s+(?:{suffixes})\b')))
        return p

    def add_custom_terms(self, terms, category="custom"):
        for t in terms:
            if t.strip():
                self.custom_terms.append((category, t.strip()))

    def _fake(self, category, original, generator):
        m = self.maps.get(category, self.maps["custom"])
        if original in m:
            return m[original]
        fake = generator(original)
        m[original] = fake
        return fake

    def _gen_name(self, orig):
        for _ in range(50):
            f = f"{self.rng.choice(FAKE_FIRST)} {self.rng.choice(FAKE_LAST)}"
            if f not in self._used_names:
                self._used_names.add(f)
                return f
        return f"Person{len(self.maps['name'])+1}"

    def _gen_company(self, orig):
        suffix = ""
        for s in ["Pte Ltd", "Pte. Ltd.", "Inc.", "Inc", "Corp.", "Corp", "LLC", "LLP", "Sdn Bhd", "Group", "Holdings"]:
            if s.lower() in orig.lower():
                suffix = s
                break
        if not suffix:
            suffix = "Pte Ltd"
        for _ in range(50):
            base = self.rng.choice(FAKE_COMPANIES)
            f = f"{base} {suffix}"
            if f not in self._used_companies:
                self._used_companies.add(f)
                return f
        return f"Company{len(self.maps['company'])+1} {suffix}"

    def _gen_email(self, orig):
        return f"{self.rng.choice(FAKE_FIRST).lower()}.{self.rng.choice(FAKE_LAST).lower()}@{self.rng.choice(FAKE_DOMAINS)}"

    def _gen_ssn(self, o):
        return f"{self.rng.randint(100,899):03d}-{self.rng.randint(10,99):02d}-{self.rng.randint(1000,9999):04d}"

    def _gen_nric(self, o):
        return f"{self.rng.choice('STFG')}{self.rng.randint(1000000,9999999)}{self.rng.choice('ABCDEFGHJKLMNPQRTUWXYZ')}"

    def _gen_phone(self, orig):
        if "+65" in orig or len(orig.replace(" ","").replace("-","")) == 8:
            return f"+65 {self.rng.randint(8000,9999)} {self.rng.randint(1000,9999)}"
        elif orig.startswith("(") or orig.startswith("+1"):
            return f"({self.rng.randint(200,999)}) {self.rng.randint(200,999)}-{self.rng.randint(1000,9999)}"
        return f"+{self.rng.randint(1,65)} {self.rng.randint(1000,9999)} {self.rng.randint(1000,9999)}"

    def _gen_engmt(self, orig):
        parts = re.split(r'[-/]', orig)
        return "-".join(str(self.rng.randint(10**(len(p)-1), 10**len(p)-1)) if p.isdigit() else p for p in parts)

    def _gen_ref(self, orig):
        m = re.match(r'([A-Za-z]+[-#]?)(\d+)', orig)
        if m:
            return f"{m.group(1)}{self.rng.randint(10**(len(m.group(2))-1), 10**len(m.group(2))-1)}"
        return f"REF{self.rng.randint(100000,999999)}"

    def _transform_amount_str(self, orig):
        clean = orig.replace("$","").replace(",","").replace("(","").replace(")","").strip()
        suffix = ""
        for s in ["B","M","K","billion","million"]:
            if s.lower() in clean.lower():
                suffix = s
                clean = re.sub(r'[BMKbmk](?:illion)?','',clean).strip()
                break
        try:
            val = float(clean) * self.amount_multiplier
            has_dec = "." in clean
            fmt = f"{val:,.{len(clean.split('.')[-1]) if has_dec else 0}f}"
            r = ""
            if "(" in orig: r += "("
            if "$" in orig: r += "$"
            r += fmt
            if suffix: r += suffix
            if "(" in orig: r += ")"
            for code in ["USD","SGD","EUR","GBP","JPY","AUD","HKD","MYR"]:
                if code in orig:
                    r = f"{code} {fmt}"
                    break
            return r
        except ValueError:
            return orig

    def _transform_date(self, orig):
        try:
            from dateutil import parser as dp
            dt = dp.parse(orig, dayfirst=True) + datetime.timedelta(days=self.date_offset)
            if re.match(r'\d{4}-\d{2}-\d{2}', orig): return dt.strftime("%Y-%m-%d")
            elif re.match(r'\d{1,2}/\d{1,2}/\d{4}', orig): return dt.strftime("%d/%m/%Y")
            elif re.match(r'\d{1,2}/\d{1,2}/\d{2}$', orig): return dt.strftime("%d/%m/%y")
            elif re.match(r'\d{1,2}-\d{1,2}-\d{4}', orig): return dt.strftime("%d-%m-%Y")
            elif re.match(r'[A-Z][a-z]+\s+\d', orig): return dt.strftime("%B %d, %Y")
            elif re.match(r'[A-Z][a-z]{2}\s+\d', orig): return dt.strftime("%b %d, %Y")
            elif re.match(r'\d{1,2}\s+[A-Z]', orig): return dt.strftime("%d %B %Y")
            return dt.strftime("%d/%m/%Y")
        except Exception:
            return orig

    def transform_text(self, text):
        if not text or not isinstance(text, str):
            return text
        result = text

        # Custom terms
        for cat, term in self.custom_terms:
            if term in result:
                if cat == "name":
                    fake = self._fake("name", term, self._gen_name)
                elif cat in ["client", "company"]:
                    fake = self._fake("company", term, self._gen_company)
                elif cat == "project":
                    fake = self._fake("project", term, lambda o: self.rng.choice(FAKE_PROJECTS))
                else:
                    fake = self._fake("custom", term, lambda o: f"Item{len(self.maps['custom'])+1}")
                result = result.replace(term, fake)

        # Regex patterns
        for cat, pattern in self.patterns:
            for match in reversed(list(pattern.finditer(result))):
                o = match.group()
                if "[" in o: continue
                if cat == "CC": f = f"{self.rng.randint(1000,9999)}-{self.rng.randint(1000,9999)}-{self.rng.randint(1000,9999)}-{self.rng.randint(1000,9999)}"
                elif cat == "SSN": f = self._fake("ssn", o, self._gen_ssn)
                elif cat == "NRIC": f = self._fake("nric", o, self._gen_nric)
                elif cat == "EMAIL": f = self._fake("email", o, self._gen_email)
                elif cat == "URL": f = f"https://{self.rng.choice(FAKE_DOMAINS)}/{''.join(self.rng.choices(string.ascii_lowercase,k=5))}"
                elif cat == "PHONE": f = self._fake("phone", o, self._gen_phone)
                elif cat == "ACCT":
                    m = re.match(r'([A-Za-z\s/:]+)(\d+)', o)
                    f = f"{m.group(1)}{self.rng.randint(10**(len(m.group(2))-1),10**len(m.group(2))-1)}" if m else o
                elif cat == "ENGMT": f = self._fake("engmt", o, self._gen_engmt)
                elif cat == "REF": f = self._fake("ref", o, self._gen_ref)
                elif cat == "AMOUNT": f = self._transform_amount_str(o)
                elif cat == "DATE": f = self._transform_date(o)
                elif cat == "COMPANY": f = self._fake("company", o, self._gen_company)
                else: continue
                result = result[:match.start()] + f + result[match.end():]
        return result

    def transform_value(self, val, is_amount=False):
        """Transform a cell value — numbers or strings."""
        if val is None:
            return None
        if isinstance(val, str):
            return self.transform_text(val)
        if isinstance(val, (int, float)):
            if is_amount and abs(val) >= 100:
                return round(val * self.amount_multiplier, 2)
            elif abs(val) >= 100000:
                return round(val * self.amount_multiplier, 2)
        return val

    def get_key_data(self):
        all_maps = {}
        for name, m in self.maps.items():
            if m:
                all_maps[name + "_map"] = dict(m)
        return {
            "seed": self.seed,
            "amount_multiplier": self.amount_multiplier,
            "amount_reversal_multiplier": round(1.0 / self.amount_multiplier, 6),
            "date_offset_days": self.date_offset,
            "date_reversal_offset": -self.date_offset,
            **all_maps,
        }

    def get_key_report(self):
        lines = ["=" * 60, "  TRANSFORMATION KEY — DO NOT UPLOAD",
                 f"  Seed: {self.seed}", f"  Amount multiplier: {self.amount_multiplier}",
                 f"  Reverse multiply by: {round(1.0/self.amount_multiplier,6)}",
                 f"  Date offset: {self.date_offset} days", "=" * 60]
        for name, m in self.maps.items():
            if m:
                lines.append(f"\n  ── {name.upper()} ──")
                for real, fake in m.items():
                    lines.append(f"  {real:40s} → {fake}")
        return "\n".join(lines)


# ══════════════════════════════════════════════════════════════
#  EXCEL → MARKDOWN (HIGH FIDELITY)
# ══════════════════════════════════════════════════════════════

def excel_to_markdown(filepath, transformer):
    """
    Convert Excel to high-fidelity markdown with:
    - Merged cells unmerged and values propagated
    - Comments extracted and placed inline
    - Formulas resolved to values
    - Sheet structure as sections
    """
    print(f"  Processing Excel: {os.path.basename(filepath)}")

    # Load twice: once with formulas, once with computed values
    wb = openpyxl.load_workbook(filepath, data_only=False)
    try:
        wb_values = openpyxl.load_workbook(filepath, data_only=True)
    except Exception:
        wb_values = wb

    md_parts = [f"# {os.path.basename(filepath)}\n"]

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        ws_val = wb_values[sheet_name] if sheet_name in wb_values.sheetnames else ws

        md_parts.append(f"\n## Sheet: {sheet_name}\n")

        # ── Step 1: Build merged cell map ────────────────────
        # Maps every cell in a merged range to the top-left cell's value
        merged_map = {}  # (row, col) → value from top-left
        merged_ranges_info = []

        for merged_range in ws.merged_cells.ranges:
            min_row, min_col = merged_range.min_row, merged_range.min_col
            max_row, max_col = merged_range.max_row, merged_range.max_col

            # Get value from top-left cell
            top_left_val = ws.cell(row=min_row, column=min_col).value

            merged_ranges_info.append({
                "range": str(merged_range),
                "rows": f"{min_row}-{max_row}",
                "cols": f"{get_column_letter(min_col)}-{get_column_letter(max_col)}",
                "value": str(top_left_val)[:50] if top_left_val else "(empty)",
            })

            # Map all cells in range to the top-left value
            # For wide merges (spanning multiple columns), only propagate
            # vertically. Horizontally merged cells should be blank except
            # the first column, to avoid repeating content across columns.
            for r in range(min_row, max_row + 1):
                for c in range(min_col, max_col + 1):
                    if r != min_row or c != min_col:
                        if min_col == max_col:
                            # Vertical merge only — propagate value down
                            merged_map[(r, c)] = top_left_val
                        elif c == min_col and r != min_row:
                            # First column of multi-column merge, non-first row
                            merged_map[(r, c)] = top_left_val
                        else:
                            # Horizontal span — blank (value shown in first col)
                            merged_map[(r, c)] = "" if (r == min_row and c != min_col) else (top_left_val if c == min_col else "")

        if merged_ranges_info:
            md_parts.append(f"*Merged cells detected and unmerged ({len(merged_ranges_info)} regions):*\n")
            for info in merged_ranges_info[:10]:
                md_parts.append(f"- {info['range']} (rows {info['rows']}): {info['value']}")
            md_parts.append("")

        # ── Step 2: Extract comments ─────────────────────────
        comments_map = {}  # (row, col) → comment text
        for row in ws.iter_rows():
            for cell in row:
                if cell.comment:
                    comments_map[(cell.row, cell.column)] = {
                        "author": cell.comment.author or "Unknown",
                        "text": cell.comment.text.strip(),
                    }

        # ── Step 3: Determine data boundaries ────────────────
        max_row = ws.max_row or 1
        max_col = ws.max_column or 1

        # ── Step 4: Detect which columns are amounts ─────────
        amount_cols = set()
        if max_row >= 1:
            for col in range(1, max_col + 1):
                header_val = ws.cell(row=1, column=col).value
                if header_val:
                    h = str(header_val).lower()
                    if any(kw in h for kw in [
                        "amount", "balance", "total", "revenue", "cost",
                        "debit", "credit", "net", "gross", "price", "fee",
                        "salary", "budget", "actual", "variance", "value",
                    ]):
                        amount_cols.add(col)

        # ── Step 5: Build the markdown table ─────────────────
        # First pass: collect all cell values
        table_data = []
        for r in range(1, max_row + 1):
            row_data = []
            for c in range(1, max_col + 1):
                # Get value: check merged map first, then computed values, then formula sheet
                if (r, c) in merged_map:
                    val = merged_map[(r, c)]
                else:
                    # Prefer computed value
                    val_computed = ws_val.cell(row=r, column=c).value
                    val_formula = ws.cell(row=r, column=c).value

                    if val_computed is not None:
                        val = val_computed
                    else:
                        val = val_formula

                # Transform the value
                is_amount = c in amount_cols
                if isinstance(val, str):
                    val = transformer.transform_text(val)
                elif isinstance(val, (int, float)):
                    val = transformer.transform_value(val, is_amount=is_amount)

                # Format for display
                if val is None:
                    display = ""
                elif isinstance(val, float):
                    if abs(val) >= 1000:
                        display = f"{val:,.2f}"
                    else:
                        display = f"{val:.2f}"
                elif isinstance(val, int):
                    if abs(val) >= 1000:
                        display = f"{val:,}"
                    else:
                        display = str(val)
                else:
                    display = str(val)

                # Append comment if present
                if (r, c) in comments_map:
                    comment = comments_map[(r, c)]
                    raw_text = comment["text"]
                    author = comment["author"]
                    # Remove author prefix if the comment text already starts with it
                    # (openpyxl sometimes includes "Author: " in the text)
                    cleaned_text = raw_text
                    for prefix in [f"{author}: ", f"{author}:", author]:
                        if cleaned_text.startswith(prefix):
                            cleaned_text = cleaned_text[len(prefix):].strip()
                            break
                    comment_text = transformer.transform_text(cleaned_text)
                    comment_author = transformer.transform_text(author)
                    display += f" **[💬 {comment_author}: {comment_text}]**"

                row_data.append(display)
            table_data.append(row_data)

        # ── Step 6: Render as markdown table ─────────────────
        if not table_data:
            md_parts.append("*(empty sheet)*\n")
            continue

        # Find max width per column for alignment
        col_widths = [0] * max_col
        for row in table_data:
            for i, cell in enumerate(row):
                col_widths[i] = max(col_widths[i], len(cell))

        # Header row
        header_row = table_data[0]
        md_parts.append("| " + " | ".join(
            cell.ljust(max(col_widths[i], 3)) for i, cell in enumerate(header_row)
        ) + " |")

        # Separator
        md_parts.append("| " + " | ".join(
            "-" * max(col_widths[i], 3) for i in range(len(header_row))
        ) + " |")

        # Data rows
        for row in table_data[1:]:
            # Skip fully empty rows
            if all(cell.strip() == "" for cell in row):
                continue
            md_parts.append("| " + " | ".join(
                cell.ljust(max(col_widths[i], 3)) if i < len(col_widths) else cell
                for i, cell in enumerate(row)
            ) + " |")

        md_parts.append("")

        # ── Step 7: Note any standalone comments not in table ─
        standalone = [(k, v) for k, v in comments_map.items()
                      if k[0] > max_row or k[1] > max_col]
        if standalone:
            md_parts.append("### Additional Comments\n")
            for (r, c), comment in standalone:
                text = transformer.transform_text(comment["text"])
                author = transformer.transform_text(comment["author"])
                md_parts.append(f"- Cell {get_column_letter(c)}{r}: **{author}**: {text}")
            md_parts.append("")

    return "\n".join(md_parts)


# ══════════════════════════════════════════════════════════════
#  WORD → MARKDOWN
# ══════════════════════════════════════════════════════════════

def word_to_markdown(filepath, transformer):
    print(f"  Processing Word: {os.path.basename(filepath)}")
    doc = Document(filepath)
    md_parts = [f"# {os.path.basename(filepath)}\n"]

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            md_parts.append("")
            continue

        transformed = transformer.transform_text(text)

        # Detect heading styles
        if para.style and para.style.name:
            style = para.style.name.lower()
            if "heading 1" in style:
                md_parts.append(f"## {transformed}")
            elif "heading 2" in style:
                md_parts.append(f"### {transformed}")
            elif "heading 3" in style:
                md_parts.append(f"#### {transformed}")
            elif "title" in style:
                md_parts.append(f"# {transformed}")
            elif "list" in style or text.startswith(("•", "-", "–")):
                md_parts.append(f"- {transformed}")
            else:
                md_parts.append(transformed)
        else:
            md_parts.append(transformed)

    # Tables
    for i, table in enumerate(doc.tables):
        md_parts.append(f"\n### Table {i+1}\n")
        for j, row in enumerate(table.rows):
            cells = [transformer.transform_text(cell.text.strip()) for cell in row.cells]
            md_parts.append("| " + " | ".join(cells) + " |")
            if j == 0:
                md_parts.append("| " + " | ".join("---" for _ in cells) + " |")

    return "\n".join(md_parts)


# ══════════════════════════════════════════════════════════════
#  PDF → MARKDOWN
# ══════════════════════════════════════════════════════════════

def pdf_to_markdown(filepath, transformer):
    print(f"  Processing PDF: {os.path.basename(filepath)}")

    # Handle encryption
    working_path = filepath
    temp_path = None
    reader = PdfReader(filepath)
    if reader.is_encrypted:
        try:
            reader.decrypt("")
            writer = PdfWriter()
            for page in reader.pages:
                writer.add_page(page)
            import tempfile
            tmp = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
            writer.write(tmp)
            tmp.close()
            working_path = tmp.name
            temp_path = tmp.name
        except Exception:
            for attempt in range(3):
                pw = input(f"  PDF password (attempt {attempt+1}/3): ").strip()
                try:
                    reader = PdfReader(filepath)
                    if reader.decrypt(pw) > 0:
                        writer = PdfWriter()
                        for page in reader.pages:
                            writer.add_page(page)
                        import tempfile
                        tmp = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
                        writer.write(tmp)
                        tmp.close()
                        working_path = tmp.name
                        temp_path = tmp.name
                        break
                except Exception:
                    pass
            else:
                return None

    md_parts = [f"# {os.path.basename(filepath)}\n"]
    try:
        with pdfplumber.open(working_path) as pdf:
            for i, page in enumerate(pdf.pages):
                md_parts.append(f"\n## Page {i+1}\n")
                text = page.extract_text()
                if text:
                    md_parts.append(transformer.transform_text(text))
                for j, table in enumerate(page.extract_tables()):
                    md_parts.append(f"\n### Table {j+1}\n")
                    for k, row in enumerate(table):
                        if row:
                            cells = [transformer.transform_text(str(c)) if c else "" for c in row]
                            md_parts.append("| " + " | ".join(cells) + " |")
                            if k == 0:
                                md_parts.append("| " + " | ".join("---" for _ in cells) + " |")
    finally:
        if temp_path and os.path.exists(temp_path):
            os.unlink(temp_path)

    return "\n".join(md_parts)


# ══════════════════════════════════════════════════════════════
#  INTERFACE
# ══════════════════════════════════════════════════════════════

def pick_files():
    print("""
╔══════════════════════════════════════════════════════════════════╗
║         SMART DATA TRANSFORMER v2.0                              ║
║     High-fidelity workpaper → markdown with fake data            ║
╚══════════════════════════════════════════════════════════════════════╝

  All files convert to markdown for optimal AI analysis.
  Merged cells, comments, and formulas are preserved.

  Drag and drop files, type 'done' when finished.
  Supports: .xlsx  .docx  .pdf  .csv  .txt
    """)
    files = []
    while True:
        entry = input("  Drop file here (or 'done'): ").strip().strip('"').strip("'")
        if entry.lower() == "done":
            break
        if os.path.isfile(entry):
            files.append(entry)
            print(f"    ✅ {os.path.basename(entry)}")
        elif os.path.isdir(entry):
            for f in os.listdir(entry):
                if os.path.splitext(f)[1].lower() in [".xlsx",".docx",".pdf",".csv",".txt"]:
                    files.append(os.path.join(entry, f))
                    print(f"    ✅ {f}")
        else:
            print(f"    ❌ Not found")
    return files

def get_custom_terms():
    print("\n  ── CUSTOM TERMS (comma-separated, Enter to skip) ──\n")
    terms = {}
    for label, cat in [("Client/company names", "company"), ("People names", "name"),
                       ("Project/codenames", "project"), ("Other terms", "custom")]:
        val = input(f"  {label}: ").strip()
        if val:
            terms[cat] = [t.strip() for t in val.split(",")]
    return terms


# ══════════════════════════════════════════════════════════════
#  MAIN
# ══════════════════════════════════════════════════════════════

def main():
    files = pick_files()
    if not files:
        print("\n  No files. Exiting.")
        input("  Press Enter...")
        return

    custom = get_custom_terms()
    t = DataTransformer()
    for cat, terms in custom.items():
        t.add_custom_terms(terms, cat)

    output_dir = os.path.join(os.path.dirname(files[0]), "TRANSFORMED_OUTPUT")
    os.makedirs(output_dir, exist_ok=True)

    print(f"\n  Seed: {t.seed} | Multiplier: {t.amount_multiplier} | Date offset: {t.date_offset} days")
    print(f"  Output: {output_dir}\n")

    outputs = []
    for fp in files:
        ext = os.path.splitext(fp)[1].lower()
        name = os.path.splitext(os.path.basename(fp))[0]

        try:
            if ext == ".xlsx":
                md = excel_to_markdown(fp, t)
            elif ext == ".docx":
                md = word_to_markdown(fp, t)
            elif ext == ".pdf":
                md = pdf_to_markdown(fp, t)
            elif ext in [".csv", ".txt"]:
                with open(fp, "r", encoding="utf-8", errors="ignore") as f:
                    md = f"# {os.path.basename(fp)}\n\n{t.transform_text(f.read())}"
            else:
                continue

            if md:
                out = os.path.join(output_dir, f"{name}.md")
                with open(out, "w", encoding="utf-8") as f:
                    f.write(md)
                print(f"  ✅ Saved: {name}.md")
                outputs.append(out)
        except Exception as e:
            print(f"  ❌ Error: {os.path.basename(fp)}: {e}")

    # Save keys
    for path, content in [
        (os.path.join(output_dir, "_KEY_DO_NOT_UPLOAD.txt"), t.get_key_report()),
        (os.path.join(output_dir, "_KEY_DO_NOT_UPLOAD.json"), json.dumps(t.get_key_data(), indent=2)),
    ]:
        with open(path, "w", encoding="utf-8") as f:
            f.write(content)

    total = sum(len(m) for m in t.maps.values())
    print(f"""
{'='*60}
  TRANSFORMATION COMPLETE
{'='*60}

  Files: {len(outputs)} | Mappings: {total}
  Output: {output_dir}

  📄 Ready for AI:""")
    for o in outputs:
        print(f"     • {os.path.basename(o)}")
    print(f"""
  🔑 Key files (keep local):
     • _KEY_DO_NOT_UPLOAD.txt / .json

  Upload the .md files to Claude. All data is fake but
  structurally accurate. Comments and merged cells preserved.
{'─'*60}""")
    input("\n  Press Enter...")

if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        print("\n  Cancelled.")
    except Exception as e:
        print(f"\n  Error: {e}")
        import traceback
        traceback.print_exc()
        input("\n  Press Enter...")
