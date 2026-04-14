"""
╔══════════════════════════════════════════════════════════════════╗
║         SMART DATA REVERSER v1.0                                 ║
║         Reverses transformed workpapers back to original         ║
║         Double-click to run — dependencies auto-install          ║
╚══════════════════════════════════════════════════════════════════════╝

Run this on your PERSONAL device.
It reads the key file and reverses all fake data back to originals.

Supports: .md  .txt  .csv  .docx  .xlsx
"""

import subprocess
import sys
import importlib

def ensure(pkg, imp=None):
    try:
        importlib.import_module(imp or pkg)
    except ImportError:
        print(f"  Installing {pkg}...")
        subprocess.check_call(
            [sys.executable, "-m", "pip", "install", pkg, "--quiet"],
            stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL
        )

print("\n  Checking dependencies...")
ensure("openpyxl")
ensure("python-docx", "docx")
print("  All dependencies ready.\n")

import os
import re
import json
import datetime

import openpyxl
from docx import Document

# ══════════════════════════════════════════════════════════════
#  REVERSER ENGINE
# ══════════════════════════════════════════════════════════════

class DataReverser:

    def __init__(self, key_data):
        self.key = key_data
        self.amount_reversal = key_data["amount_reversal_multiplier"]
        self.date_reversal = key_data["date_reversal_offset"]

        # Build reverse mappings (fake → real)
        self.reverse_maps = {}
        for map_name in ["name_map", "company_map", "email_map", "ssn_map",
                         "nric_map", "phone_map", "engmt_map", "ref_map",
                         "project_map", "custom_map"]:
            if map_name in key_data and key_data[map_name]:
                for real, fake in key_data[map_name].items():
                    self.reverse_maps[fake] = real

        # Sort by length descending so longer matches replace first
        self.sorted_fakes = sorted(self.reverse_maps.keys(), key=len, reverse=True)

        print(f"  Loaded key: seed {key_data['seed']}")
        print(f"  Amount reversal multiplier: {self.amount_reversal}")
        print(f"  Date reversal offset: {self.date_reversal} days")
        print(f"  Total mappings to reverse: {len(self.reverse_maps)}")

    def reverse_text(self, text):
        """Reverse all fake data in text back to originals."""
        if not text or not isinstance(text, str):
            return text

        result = text

        # Step 1: Replace fake names/companies/emails etc. with originals
        for fake in self.sorted_fakes:
            if fake in result:
                result = result.replace(fake, self.reverse_maps[fake])

        # Step 2: Reverse dollar amounts
        # Find amounts and reverse the multiplier
        amount_patterns = [
            re.compile(r'[\(]?\$\s?\d{1,3}(?:,\d{3})*(?:\.\d{1,2})?(?:\s?[MBKmb](?:illion)?)?\)?'),
            re.compile(r'\b(?:USD|SGD|EUR|GBP|JPY|AUD|HKD|MYR)\s?\d{1,3}(?:,\d{3})*(?:\.\d{1,2})?\b'),
        ]

        for pattern in amount_patterns:
            matches = list(pattern.finditer(result))
            for match in reversed(matches):
                original_str = match.group()
                reversed_str = self._reverse_amount(original_str)
                if reversed_str != original_str:
                    result = result[:match.start()] + reversed_str + result[match.end():]

        # Step 3: Reverse dates
        date_pattern = re.compile(
            r'\b(?:'
            r'\d{1,2}[/-]\d{1,2}[/-]\d{2,4}'
            r'|(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{1,2},?\s+\d{4}'
            r'|\d{1,2}\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{4}'
            r'|\d{4}-\d{2}-\d{2}'
            r')\b',
            re.IGNORECASE
        )
        matches = list(date_pattern.finditer(result))
        for match in reversed(matches):
            original_str = match.group()
            reversed_str = self._reverse_date(original_str)
            if reversed_str != original_str:
                result = result[:match.start()] + reversed_str + result[match.end():]

        return result

    def _reverse_amount(self, amount_str):
        """Reverse a dollar amount by applying reversal multiplier."""
        clean = amount_str.replace("$", "").replace(",", "").replace("(", "").replace(")", "").strip()

        suffix = ""
        for s in ["B", "M", "K", "billion", "million"]:
            if s.lower() in clean.lower():
                suffix = s
                clean = re.sub(r'[BMKbmk](?:illion)?', '', clean).strip()
                break

        try:
            value = float(clean)
            original_value = value * self.amount_reversal

            is_parens = "(" in amount_str
            has_dollar = "$" in amount_str

            if "." in clean:
                decimals = len(clean.split(".")[-1])
                formatted = f"{original_value:,.{decimals}f}"
            else:
                formatted = f"{original_value:,.0f}"

            result = ""
            if is_parens:
                result += "("
            if has_dollar:
                result += "$"
            result += formatted
            if suffix:
                result += suffix
            if is_parens:
                result += ")"

            for code in ["USD", "SGD", "EUR", "GBP", "JPY", "AUD", "HKD", "MYR"]:
                if code in amount_str:
                    result = f"{code} {formatted}"
                    break

            return result
        except ValueError:
            return amount_str

    def _reverse_date(self, date_str):
        """Reverse a date by applying reversal offset."""
        try:
            from dateutil import parser as dateparser
        except ImportError:
            return date_str

        try:
            dt = dateparser.parse(date_str, dayfirst=False)
            original_dt = dt + datetime.timedelta(days=self.date_reversal)

            if re.match(r'\d{4}-\d{2}-\d{2}', date_str):
                return original_dt.strftime("%Y-%m-%d")
            elif re.match(r'\d{1,2}/\d{1,2}/\d{4}', date_str):
                return original_dt.strftime("%m/%d/%Y")
            elif re.match(r'\d{1,2}/\d{1,2}/\d{2}$', date_str):
                return original_dt.strftime("%m/%d/%y")
            elif re.match(r'\d{1,2}-\d{1,2}-\d{4}', date_str):
                return original_dt.strftime("%m-%d-%Y")
            elif re.match(r'[A-Z][a-z]+\s+\d', date_str):
                return original_dt.strftime("%B %d, %Y")
            elif re.match(r'[A-Z][a-z]{2}\s+\d', date_str):
                return original_dt.strftime("%b %d, %Y")
            elif re.match(r'\d{1,2}\s+[A-Z]', date_str):
                return original_dt.strftime("%d %B %Y")
            else:
                return original_dt.strftime("%Y-%m-%d")
        except Exception:
            return date_str

    def reverse_number(self, value):
        """Reverse a numeric value."""
        if isinstance(value, str):
            # Check if it's a placeholder-style amount
            return self.reverse_text(value)
        if isinstance(value, (int, float)):
            if abs(value) >= 100:
                return round(value * self.amount_reversal, 2)
        return value


# ══════════════════════════════════════════════════════════════
#  FILE PROCESSORS
# ══════════════════════════════════════════════════════════════

def reverse_text_file(filepath, reverser, output_path):
    """Reverse a .txt, .md, or .csv file."""
    print(f"  Reversing: {os.path.basename(filepath)}")
    with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
        content = f.read()
    reversed_content = reverser.reverse_text(content)
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(reversed_content)
    print(f"  ✅ Saved: {output_path}")


def reverse_excel(filepath, reverser, output_path):
    """Reverse an Excel workbook."""
    print(f"  Reversing Excel: {os.path.basename(filepath)}")
    wb = openpyxl.load_workbook(filepath, data_only=False)

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        print(f"    Sheet: {sheet_name}")

        for row in ws.iter_rows():
            for cell in row:
                if cell.value is not None:
                    if isinstance(cell.value, str):
                        cell.value = reverser.reverse_text(cell.value)
                    elif isinstance(cell.value, (int, float)):
                        cell.value = reverser.reverse_number(cell.value)

    wb.save(output_path)
    print(f"  ✅ Saved: {output_path}")


def reverse_word(filepath, reverser, output_path):
    """Reverse a Word document."""
    print(f"  Reversing Word: {os.path.basename(filepath)}")
    doc = Document(filepath)
    count = 0

    for para in doc.paragraphs:
        for run in para.runs:
            if run.text:
                new_text = reverser.reverse_text(run.text)
                if new_text != run.text:
                    run.text = new_text
                    count += 1

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.text:
                            new_text = reverser.reverse_text(run.text)
                            if new_text != run.text:
                                run.text = new_text
                                count += 1

    for section in doc.sections:
        for hf in [section.header, section.footer]:
            if hf:
                for para in hf.paragraphs:
                    for run in para.runs:
                        if run.text:
                            new_text = reverser.reverse_text(run.text)
                            if new_text != run.text:
                                run.text = new_text
                                count += 1

    doc.save(output_path)
    print(f"  ✅ Saved: {output_path} ({count} runs reversed)")


# ══════════════════════════════════════════════════════════════
#  MAIN
# ══════════════════════════════════════════════════════════════

def main():
    print("""
╔══════════════════════════════════════════════════════════════════╗
║              SMART DATA REVERSER v1.0                             ║
║         Reverse transformed files back to originals              ║
╚══════════════════════════════════════════════════════════════════════╝

  Run this on your personal device.
  You need: the transformed file(s) + the key file.
    """)

    # Step 1: Load key file
    print("  STEP 1: Load the key file\n")
    key_path = input("  Drop _KEY_DO_NOT_UPLOAD.json here: ").strip().strip('"').strip("'")

    if not os.path.isfile(key_path):
        print(f"  ❌ Key file not found: {key_path}")
        input("  Press Enter to close...")
        return

    with open(key_path, "r", encoding="utf-8") as f:
        key_data = json.load(f)

    reverser = DataReverser(key_data)

    # Step 2: Load transformed files
    print("\n  STEP 2: Load transformed files\n")
    files = []
    while True:
        entry = input("  Drop file here (or 'done'): ").strip().strip('"').strip("'")
        if entry.lower() == "done":
            break
        if os.path.isfile(entry):
            files.append(entry)
            print(f"    ✅ {os.path.basename(entry)}")
        else:
            print(f"    ❌ Not found")

    if not files:
        print("\n  No files. Exiting.")
        input("  Press Enter to close...")
        return

    # Step 3: Reverse
    output_dir = os.path.join(os.path.dirname(files[0]), "REVERSED_ORIGINALS")
    os.makedirs(output_dir, exist_ok=True)

    print(f"\n  Output: {output_dir}\n")

    output_files = []
    for filepath in files:
        ext = os.path.splitext(filepath)[1].lower()
        name = os.path.splitext(os.path.basename(filepath))[0]

        try:
            if ext in [".md", ".txt", ".csv"]:
                out = os.path.join(output_dir, f"{name}_ORIGINAL{ext}")
                reverse_text_file(filepath, reverser, out)
                output_files.append(out)

            elif ext == ".xlsx":
                out = os.path.join(output_dir, f"{name}_ORIGINAL.xlsx")
                reverse_excel(filepath, reverser, out)
                output_files.append(out)

            elif ext == ".docx":
                out = os.path.join(output_dir, f"{name}_ORIGINAL.docx")
                reverse_word(filepath, reverser, out)
                output_files.append(out)

            else:
                print(f"  ⚠️  Unsupported format: {ext} — skipping")

        except Exception as e:
            print(f"  ❌ Error: {os.path.basename(filepath)}: {e}")

    # Summary
    print(f"""
{'='*60}
  REVERSAL COMPLETE
{'='*60}

  Files reversed: {len(output_files)}
  Output folder:  {output_dir}

  📄 RESTORED FILES:""")
    for f in output_files:
        print(f"     • {os.path.basename(f)}")

    print(f"""
{'─'*60}

  These files now contain the original data.
  You can upload them to Claude for analysis.

  ⚠️  Remember to delete these files from your personal
     device when you're done with the review.
{'─'*60}
    """)

    input("  Press Enter to close...")


if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        print("\n\n  Cancelled.")
    except Exception as e:
        print(f"\n  Error: {e}")
        import traceback
        traceback.print_exc()
        input("\n  Press Enter to close...")
